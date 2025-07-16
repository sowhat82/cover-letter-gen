
import streamlit as st
import fitz  # PyMuPDF
import requests
import zipfile
import io
from docx import Document as DocxDocument
from docx.shared import Pt
#from streamlit_analytics import track
import os
import re

def extract_text_from_jd(file):
    if file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.name.endswith(".pdf"):
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in pdf])
    elif file.name.endswith(".docx"):
        doc = DocxDocument(file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return ""

# --- CONFIG ---
st.set_page_config(page_title="AI Cover Letter Generator", layout="centered")

# --- TRACK USAGE ---
#with track():

st.title("📄 GPT-Powered Cover Letter Generator")

st.info("🔒 Your files are processed securely and never stored. All data is handled temporarily during your session and is erased when you leave the page.")

st.markdown("""
Upload your resume and one or more job descriptions.
The app will generate tailored cover letters for each JD.
""")

# 🔑 Securely load your OpenRouter API key from Streamlit secrets
api_key = st.secrets.get("openrouter_api_key", "")


if not api_key.strip():
    st.error("❌ OpenRouter API key not found. Please set it in Streamlit secrets.")
    st.stop()

headers = {
    "Authorization": f"Bearer {api_key}",
    "Content-Type": "application/json"
}
url = "https://openrouter.ai/api/v1/chat/completions"

# --- Uploads ---
resume_file = st.file_uploader("📎 Upload your resume (PDF)", type=["pdf"])
jd_files = st.file_uploader(
    "📄 Upload one or more Job Descriptions (PDF, Word, or TXT):",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True
)

# Accept text input for job descriptions
jd_text_input = st.text_area(
    "📋 OR  Paste one or more job descriptions here",
    help="To separate multiple JDs, insert a new line with three hyphens like this:\n\n\"---\"\n\nAvoid using --- inside the job descriptions themselves."
)


# --- Helper Functions ---
def test_openrouter_key():
    test_prompt = {
        "model": "openai/gpt-3.5-turbo",
        "messages": [{"role": "user", "content": "Say hello"}]
    }

    response = requests.post(url, headers=headers, json=test_prompt)
    st.write("Test API status:", response.status_code)
    st.write("Test response:", response.text)

#if st.button("🔍 Test OpenRouter"):
#   test_openrouter_key()


def extract_text_from_pdf(uploaded_pdf):
    doc = fitz.open(stream=uploaded_pdf.read(), filetype="pdf")
    return "".join([page.get_text() for page in doc])

def generate_filename_from_jd(jd_text):
    naming_prompt = {
        "model": "mistralai/mistral-7b-instruct",
        "messages": [
            {
                "role": "system",
                "content": "Summarize the job description into a 3–5 word job title (no company name, no punctuation). Output only the job title."
            },
            {
                "role": "user",
                "content": jd_text
            }
        ]
    }

    try:
        response = requests.post(url, headers=headers, json=naming_prompt)
        title = response.json()["choices"][0]["message"]["content"].strip()
        words = title.split()
        filtered = [w for w in words if w.lower() not in ['collaborative', 'analytical', 'driven', 'stakeholder', 'leader']]
        short_name = "_".join(filtered[:4])
        return re.sub(r'[^A-Za-z0-9_]', '', short_name) or "UnnamedRole"
    except:
        return "UnnamedRole"

def generate_cover_letter(resume, jd, length_option, custom_instructions):

    if "Short" in length_option:
        word_target = "around 200 words"
    elif "Long" in length_option:
        word_target = "at least 400 words"
    else:
        word_target = "around 300 words"

    if custom_instructions.strip():
        user_input = (
            f"My resume: {resume}\n\n"
            f"Job description: {jd}\n\n"
            f"Please include the following personal context in the cover letter: {custom_instructions.strip()}"
        )
    else:
        user_input = f"My resume: {resume}\n\nJob description: {jd}"

    prompt = {
        "model": "mistralai/mistral-7b-instruct",
        "messages": [
            {
                "role": "system",
                "content": (
                    f"You are a career assistant. Write a tailored, confident cover letter based on the user's resume and job description. "
                    f"The letter should be {word_target}. "
                    f"If the user provides any personal notes or context (such as a referral, connection, or why they’re interested), weave that naturally into the letter. "
                    f"Do not list these as standalone bullet points. Start directly with the greeting (e.g. 'Dear Hiring Manager') and do not include any names, addresses, phone numbers, or emails. "
                    f"Sign off with 'Your Name'."
                )
            },
            {
                "role": "user",
                "content": user_input
            }
        ]
    }

    response = requests.post(url, headers=headers, json=prompt)

    # 🔍 Debug: print response details
#    st.write("API status:", response.status_code)

    try:
        result_json = response.json()
        if "choices" in result_json:
            return result_json["choices"][0]["message"]["content"]
        else:
            st.error("❌ API response does not contain expected content.")
            st.write("Full response:", result_json)
            st.stop()
    except Exception as e:
        st.error("❌ Failed to process the API response.")
        st.write("Exception details:", str(e))
        st.write("Raw response:", response.text)
        st.stop()

    return None

# --- Cover Letter Length Option ---
length_option = st.selectbox(
    "📝 Choose cover letter length:",
    ["Short (≈ 200 words)", "Medium (≈ 300 words)", "Long (≈ 400+ words)"],
    index=1
)

# --- Optional Custom Instructions ---
custom_instructions = st.text_area(
    "🧠 Optional: Add custom notes (e.g. referral, where you found the job, specific interest)",
    placeholder="Example: I was referred by a former colleague, John Tan. I’ve also followed your company’s sustainability work since 2021.",
    height=100
)

# --- Main Logic ---
if st.button("✍️ Generate Cover Letters"):
    if not resume_file or (not jd_files and not jd_text_input.strip()):
        st.warning("Please upload your resume and provide at least one job description.")
        st.stop()

    resume_text = extract_text_from_pdf(resume_file)

    # ✅ Gather job descriptions from files and text input
    jd_texts = []

    if jd_files:
        for file in jd_files:
            jd_texts.append(extract_text_from_jd(file))


    if jd_text_input.strip():
        pasted_jds = [jd.strip() for jd in jd_text_input.split('---') if jd.strip()]
        jd_texts.extend(pasted_jds)

    with st.spinner("✍️ Generating your cover letters... This may take 15–30 seconds."):

        output_zip = io.BytesIO()
        with zipfile.ZipFile(output_zip, "w") as zipf:
            for i, jd_text in enumerate(jd_texts, 1):
                slug = generate_filename_from_jd(jd_text)
                docx_filename = f"CoverLetter_{slug}_{i}.docx"

                cover_letter = generate_cover_letter(resume_text, jd_text, length_option, custom_instructions)

                doc_stream = io.BytesIO()
                doc = DocxDocument()
                for para in cover_letter.strip().split("\n\n"):
                    p = doc.add_paragraph(para)
                    p.style.font.name = 'Calibri'
                    p.style.font.size = Pt(11)
                doc.save(doc_stream)
                doc_stream.seek(0)

                zipf.writestr(docx_filename, doc_stream.read())

    st.success("✅ Cover letters generated!")
    st.download_button("📥 Download All Cover Letters (ZIP)", data=output_zip.getvalue(), file_name="CoverLetters.zip", mime="application/zip")

st.caption("⚠️ AI-generated. Please review and personalize before sending.")

st.caption("⚠️ Heads up: This tool runs on a free-tier server and shared AI API, so responses may take a while—or may temporarily fail if traffic is high.")

st.markdown("☕ If this tool saved you time, [you can support me here](https://buymeacoffee.com/alvincheong). It helps me keep building more useful AI tools!")
